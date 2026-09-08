#!/usr/bin/env python3
"""
k3-doc-intel — Document Intelligence MCP Server
Ingests, extracts, chunks, and semantically searches PDFs, DOCX, and text documents
using PyMuPDF, python-docx, and Gemini MRL embeddings.
Supports isolated per-project document indices (.agents/doc_index.pkl).
"""

import os
import pickle
import hashlib
import time
from pathlib import Path
from typing import Optional, List, Dict, Any
import numpy as np
import requests
import fitz  # PyMuPDF
import docx  # python-docx

from mcp.server.fastmcp import FastMCP

mcp = FastMCP("k3-doc-intel")

DEFAULT_GLOBAL_INDEX = Path(
    os.environ.get(
        "K3_DOC_INDEX",
        Path(__file__).resolve().parent.parent / "src" / "doc_index.pkl"
    )
).resolve()
TARGET_DIMENSION = 768
MODEL = "models/gemini-embedding-001"


def _get_api_key() -> str:
    key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not key:
        raise ValueError("Missing GEMINI_API_KEY or GOOGLE_API_KEY.")
    return key


def _embed_single(text: str) -> np.ndarray:
    """Embeds text using gemini-embedding-001 (768 MRL dimensions, L2 normalized)."""
    api_key = _get_api_key()
    url = f"https://generativelanguage.googleapis.com/v1beta/{MODEL}:embedContent"
    headers = {
        "Content-Type": "application/json",
        "x-goog-api-key": api_key,
    }
    payload = {
        "model": MODEL,
        "content": {"parts": [{"text": text[:8000]}]},
    }
    r = requests.post(url, headers=headers, json=payload, timeout=20)
    r.raise_for_status()
    raw = r.json()["embedding"]["values"]

    vec = np.array(raw, dtype=np.float32)
    if TARGET_DIMENSION < vec.shape[0]:
        vec = vec[:TARGET_DIMENSION]
    norm = np.linalg.norm(vec)
    if norm > 0:
        vec = vec / norm
    return vec


def _resolve_index_path(project_dir: Optional[str] = None) -> Path:
    """Resolves target doc_index.pkl path without disk side-effects."""
    if project_dir:
        p = Path(project_dir).resolve()
        return p / ".agents" / "doc_index.pkl"
    return DEFAULT_GLOBAL_INDEX


def _extract_text_from_file(file_path: Path) -> str:
    """Extracts raw text from PDF, DOCX, or text/code formats."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        doc = fitz.open(str(file_path))
        pages = []
        for i, page in enumerate(doc):
            t = page.get_text().strip()
            if t:
                pages.append(f"--- [Page {i + 1}] ---\n{t}")
        doc.close()
        return "\n\n".join(pages)

    elif suffix in (".docx", ".doc"):
        doc = docx.Document(str(file_path))
        paras = []
        for p in doc.paragraphs:
            if p.text.strip():
                paras.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paras.append(row_text)
        return "\n\n".join(paras)

    else:
        # Fallback to UTF-8 plain text read
        return file_path.read_text(encoding="utf-8", errors="replace")


def _chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> List[str]:
    """Splits text into chunks respecting paragraph and line boundaries."""
    chunk_size = max(50, chunk_size)
    overlap = max(0, min(overlap, chunk_size // 2))

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            # Look backwards for newline or space
            split_at = text.rfind("\n", start + overlap, end)
            if split_at == -1:
                split_at = text.rfind(" ", start + overlap, end)
            if split_at != -1:
                end = split_at

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = max(end - overlap, start + 1)
        if end >= len(text):
            break

    return chunks


@mcp.tool()
def doc_extract_text(file_path: str, max_chars: int = 50000) -> str:
    """
    Extracts structured plain text from PDF, DOCX, Markdown, or text files.
    Preserves page breaks and table cell structure.
    Args:
        file_path: Absolute path to the document file.
        max_chars: Maximum characters to return (default: 50,000).
    """
    p = Path(file_path).resolve()
    if not p.exists():
        return f"Error: File not found at {p}"

    try:
        text = _extract_text_from_file(p)
        truncated = len(text) > max_chars
        content = text[:max_chars]
        hdr = f"=== Extracted Document: {p.name} ({len(text):,} chars total) ===\n"
        if truncated:
            hdr += f"[Notice: Showing first {max_chars:,} characters]\n"
        return hdr + "\n" + content
    except Exception as e:
        return f"Extraction error: {e}"


@mcp.tool()
def doc_chunk_and_embed(
    file_path: str,
    project_dir: Optional[str] = None,
    chunk_size: int = 2000
) -> str:
    """
    Chunks a document (PDF, DOCX, TXT) and generates Gemini MRL embeddings.
    Stores chunks and vectors incrementally in the target project's doc_index.pkl.
    Args:
        file_path: Absolute path to the document to ingest.
        project_dir: Optional project directory (defaults to global doc_index.pkl).
        chunk_size: Target characters per chunk (default: 2,000).
    """
    p = Path(file_path).resolve()
    if not p.exists():
        return f"Error: File not found at {p}"

    idx_path = _resolve_index_path(project_dir)
    idx_path.parent.mkdir(parents=True, exist_ok=True)

    # Load existing index
    index = {}
    if idx_path.exists():
        try:
            with open(idx_path, "rb") as f:
                index = pickle.load(f)
        except Exception:
            index = {}

    # Extract text & calculate file hash
    try:
        raw_text = _extract_text_from_file(p)
        file_hash = hashlib.md5(raw_text.encode("utf-8", errors="ignore")).hexdigest()
    except Exception as e:
        return f"Failed to extract document: {e}"

    # Check if already embedded with same hash
    existing_for_file = [k for k, v in index.items() if v.get("source_file") == str(p)]
    if existing_for_file and index[existing_for_file[0]].get("file_hash") == file_hash:
        return f"Already Up-To-Date: {p.name} ({len(existing_for_file)} chunks already indexed with matching MD5 {file_hash[:8]})."

    # Remove old chunks for this file if hash changed
    for k in existing_for_file:
        del index[k]

    chunks = _chunk_text(raw_text, chunk_size=chunk_size)
    if not chunks:
        return f"Warning: Document {p.name} yielded 0 non-empty chunks."

    embedded_count = 0
    t0 = time.time()
    for i, chunk in enumerate(chunks):
        chunk_key = f"{p.name}::chunk[{i}]"
        try:
            vec = _embed_single(chunk)
            index[chunk_key] = {
                "source_file": str(p),
                "file_hash": file_hash,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "vector": vec,
                "snippet": chunk[:300],
                "text": chunk,
                "ingested_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            embedded_count += 1
            time.sleep(0.1)  # Gentle rate limit
        except Exception as e:
            return f"Error embedding chunk {i} of {p.name}: {e}"

    # Save index
    idx_path.parent.mkdir(parents=True, exist_ok=True)
    with open(idx_path, "wb") as f:
        pickle.dump(index, f)

    elapsed = time.time() - t0
    return (
        f"Successfully Ingested: {p.name}\n"
        f"  - Chunks Created : {embedded_count} / {len(chunks)}\n"
        f"  - Ingestion Time : {elapsed:.2f}s\n"
        f"  - Index Location : {idx_path}\n"
        f"  - Total In Index : {len(index)} chunks across all docs"
    )


@mcp.tool()
def doc_search(
    query: str,
    top_k: int = 5,
    project_dir: Optional[str] = None
) -> str:
    """
    Performs semantic vector search across ingested documents in the project's doc_index.pkl.
    Args:
        query: Question or search phrase (e.g. 'indemnification cap tipping basket', 'escrow period').
        top_k: Number of relevant chunks to retrieve (default: 5).
        project_dir: Optional project directory to search within.
    """
    if not query or not query.strip():
        return "Error: Search query cannot be empty."

    idx_path = _resolve_index_path(project_dir)
    if not idx_path.exists():
        return f"No document index found at {idx_path}. Use doc_chunk_and_embed first."

    try:
        with open(idx_path, "rb") as f:
            index = pickle.load(f)
    except Exception as e:
        return f"Failed to load document index: {e}"

    if not index:
        return f"Document index at {idx_path} is empty."

    try:
        q_vec = _embed_single(query)
    except Exception as e:
        return f"Failed to embed query: {e}"

    top_k = min(max(1, top_k), 50)

    try:
        keys = list(index.keys())
        matrix = np.vstack([index[k]["vector"] for k in keys])
        scores = np.dot(matrix, q_vec)

        actual_k = min(top_k, len(keys))
        partition_indices = np.argpartition(scores, -actual_k)[-actual_k:]
        sorted_indices = partition_indices[np.argsort(-scores[partition_indices])]
    except Exception as e:
        return f"Failed to compute document similarity: {e}"

    results = []
    for rank, idx in enumerate(sorted_indices, 1):
        score = float(scores[idx])
        key = keys[idx]
        item = index[key]
        src_name = Path(item.get("source_file", key)).name
        chunk_idx = item.get("chunk_index", 0)
        tot = item.get("total_chunks", 1)
        text_preview = item.get("text", item.get("snippet", ""))[:400].replace("\n", " ")

        results.append(
            f"{rank}. [Score: {score:.4f}] {src_name} (Chunk {chunk_idx + 1}/{tot})\n"
            f"   Text: {text_preview}..."
        )

    return (
        f"=== Document Search Results ({len(results)} matches for '{query}') ===\n"
        f"Index: {idx_path}\n\n" + "\n\n".join(results)
    )


@mcp.tool()
def doc_list_ingested(project_dir: Optional[str] = None) -> str:
    """
    Lists all documents currently indexed in the target doc_index.pkl,
    showing chunk counts, file size, and ingestion timestamps.
    """
    idx_path = _resolve_index_path(project_dir)
    if not idx_path.exists():
        return f"No document index found at {idx_path}."

    try:
        with open(idx_path, "rb") as f:
            index = pickle.load(f)
    except Exception as e:
        return f"Failed to load document index: {e}"

    docs: Dict[str, Dict[str, Any]] = {}
    for item in index.values():
        src = item.get("source_file", "unknown")
        if src not in docs:
            docs[src] = {
                "chunks": 0,
                "ingested_at": item.get("ingested_at", "unknown"),
                "hash": item.get("file_hash", "")[:8],
            }
        docs[src]["chunks"] += 1

    lines = [
        f"=== Ingested Document Index ===",
        f"Index Location: {idx_path}",
        f"Total Ingested Documents: {len(docs)}",
        f"Total Stored Chunks: {len(index):,}\n",
    ]

    for src, meta in sorted(docs.items(), key=lambda x: x[0]):
        p = Path(src)
        sz = f"{p.stat().st_size / 1024:.1f} KB" if p.exists() else "Missing on disk"
        lines.append(
            f"- {p.name:<35} | {meta['chunks']:>3} chunks | {sz:<10} | Ingested: {meta['ingested_at']}"
        )

    return "\n".join(lines)


if __name__ == "__main__":
    mcp.run()
